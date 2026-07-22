from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path("/workspace/scratch/6f657a5e2035")
REPO = ROOT / "amr-changing-world"
AUDIT = REPO / "outputs" / "phase4_analysis"
OUT = ROOT / "outputs" / "phase4"
FIG = OUT / "figures"
OUT.mkdir(parents=True, exist_ok=True)
FIG.mkdir(parents=True, exist_ok=True)

main = pd.read_csv(AUDIT / "main_conflict_models.csv")
diagnostics = pd.read_csv(AUDIT / "model_diagnostics.csv")
sensitivity = pd.read_csv(AUDIT / "sensitivity_models.csv")
standard = pd.read_csv(AUDIT / "standardised_amr.csv")
trends = pd.read_csv(AUDIT / "standardised_trends.csv")
calibration = pd.read_csv(AUDIT / "standardisation_calibration.csv")
std_diag = pd.read_csv(AUDIT / "standardisation_diagnostics.csv")
mic = pd.read_csv(AUDIT / "mic_models.csv")
events = pd.read_csv(AUDIT / "exploratory_event_trajectories.csv")

labels = {
    "ECO_CAZ_R": r"$\it{E.\ coli}$–ceftazidime",
    "KPN_CAZ_R": r"$\it{K.\ pneumoniae}$–ceftazidime",
    "KPN_MEM_R": r"$\it{K.\ pneumoniae}$–meropenem",
    "ABA_MEM_R": r"$\it{A.\ baumannii}$–meropenem",
}

plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 9})
fig, ax = plt.subplots(figsize=(7.1, 3.3))
ordered = main.iloc[::-1].reset_index(drop=True)
y = np.arange(len(ordered))
ax.errorbar(ordered.or_per_doubling_1plus_events, y,
            xerr=[ordered.or_per_doubling_1plus_events - ordered.or_ci_low,
                  ordered.or_ci_high - ordered.or_per_doubling_1plus_events],
            fmt="o", color="black", ecolor="#555555", capsize=3, markersize=5)
ax.axvline(1, color="#777777", linestyle="--", linewidth=1)
ax.set_yticks(y, [labels[x] for x in ordered.endpoint_id])
ax.set_xlabel("Odds ratio per doubling of 1 + preceding-year conflict events (95% CI)")
ax.set_xlim(0.85, 1.25)
ax.grid(axis="x", color="#DDDDDD", linewidth=.6)
for spine in ["top", "right"]:
    ax.spines[spine].set_visible(False)
fig.tight_layout()
forest_path = FIG / "Figure_1_Primary_Conflict_Effects.png"
fig.savefig(forest_path, dpi=300, bbox_inches="tight")
plt.close(fig)

fig, axes = plt.subplots(2, 2, figsize=(7.2, 5.5), sharex=True)
for ax, endpoint in zip(axes.flat, labels):
    part = trends.loc[trends.endpoint_id.eq(endpoint)].sort_values("year")
    ax.plot(part.year, part.isolate_weighted_standardised_pct, color="black", marker="o", label="Standardised")
    ax.plot(part.year, part.pooled_crude_resistance_pct, color="#888888", marker="s", linestyle="--", label="Crude")
    ax.set_title(labels[endpoint], fontsize=9)
    ax.set_ylabel("Resistance (%)")
    ax.grid(axis="y", color="#E5E5E5", linewidth=.5)
    ax.spines[["top", "right"]].set_visible(False)
axes[1, 0].set_xlabel("Outcome year")
axes[1, 1].set_xlabel("Outcome year")
axes[0, 0].legend(frameon=False, fontsize=8)
fig.tight_layout()
trend_path = FIG / "Figure_2_Crude_and_Standardised_Trends.png"
fig.savefig(trend_path, dpi=300, bbox_inches="tight")
plt.close(fig)

NAVY = "275D7A"
PALE_BLUE = "DDEBF3"
PALE_GRAY = "F2F4F5"
AMBER = "FCE8B2"
GREEN = "D9EAD3"
BORDER = "CCD4D9"
DARK = RGBColor(36, 50, 63)
MUTED = RGBColor(102, 114, 124)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)


def set_cell(cell, value, bold=False, color=None, size=8.1, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(value))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    cell_border(cell)


def set_table_geometry(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(sum(widths_cm) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        grid.append(col)
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            dxa = int(width / 2.54 * 1440)
            row.cells[index].width = Cm(width)
            tc_w = row.cells[index]._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, size=8.1):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, text in enumerate(headers):
        set_cell(table.rows[0].cells[i], text, bold=True, color="FFFFFF", size=size, center=True)
        shade(table.rows[0].cells[i], NAVY)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell(cells[i], value, size=size, center=i > 0)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_note(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=2)
    set_cell(table.cell(0, 0), label, bold=True, color=NAVY, size=9)
    set_cell(table.cell(0, 1), text, size=9)
    shade(table.cell(0, 0), fill)
    shade(table.cell(0, 1), fill)
    set_table_geometry(table, [3.0, 13.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(.492)
section.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, before, after in [("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 12, 8, 4)]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(NAVY)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

# Memo masthead: standard_business_brief preset.
header = section.header.paragraphs[0]
header.text = "AMR IN A CHANGING WORLD  |  PHASE 4"
header.runs[0].font.name = "Calibri"
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = MUTED
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("Statistical analysis report  |  23 July 2026").font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("PHASE 4 STATISTICAL ANALYSIS")
r.bold = True
r.font.name = "Calibri"
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("Standardised AMR estimates and confirmatory conflict models")
r.font.size = Pt(14)
r.font.color.rgb = MUTED
for label, value in [
    ("Project", "Human AMR trends in a changing world"),
    ("Protocol", "Version 1.0; Phase 3 sample lock retained"),
    ("Pipeline", "Version 0.4.0"),
    ("Analysis date", "23 July 2026"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)
add_note(doc, "Primary conclusion", "Within the locked 2019–2024 ATLAS sample, preceding-year political-violence event intensity was not associated with any of the four prespecified resistance outcomes after multiplicity correction. Confidence intervals remain compatible with modest decreases or increases, so the result should be interpreted as no detectable association, not proof of no effect.", GREEN)

doc.add_heading("1. Confirmatory results", level=1)
doc.add_paragraph(
    "Each endpoint was analysed separately using isolate-level logistic regression with country and calendar-year fixed effects, adjustment for age group, sex, specimen source, clinical specialty and lagged annual temperature anomaly, and country-clustered inference. One exposure unit is a doubling of 1 + preceding-year ACLED political-violence events. Intermediate isolates were classified as not resistant, consistent with the prespecified R-versus-non-R outcome."
)
rows = []
for row in main.itertuples(index=False):
    rows.append([
        row.endpoint_id, f"{row.tested_isolates:,}", row.countries,
        f"{row.or_per_doubling_1plus_events:.3f} ({row.or_ci_low:.3f}–{row.or_ci_high:.3f})",
        f"{row.cluster_t_p:.3f}", f"{row.holm_p:.3f}",
        f"{100*row.average_marginal_risk_difference:+.2f} ({100*row.amrd_ci_low:+.2f} to {100*row.amrd_ci_high:+.2f})",
    ])
add_table(doc, ["Endpoint", "Isolates", "Countries", "Adjusted OR (95% CI)", "p value", "Holm p", "Marginal change, pp (95% CI)"],
          rows, [2.2, 1.8, 1.6, 3.2, 1.5, 1.5, 4.7], size=7.8)
doc.add_picture(str(forest_path), width=Inches(6.25))
p = doc.add_paragraph("Figure 1. Adjusted conflict-effect estimates. The vertical line denotes no association (OR = 1).")
p.style = doc.styles["Caption"]

doc.add_heading("Interpretation", level=2)
add_bullets(doc, [
    "No primary endpoint met the Holm-adjusted 0.05 threshold; adjusted p values ranged from 0.776 to 0.968.",
    "Point estimates were slightly below 1 for ceftazidime resistance and slightly above 1 for meropenem resistance, with every confidence interval crossing 1.",
    "Average marginal changes ranged from −0.76 to +0.64 percentage points per exposure doubling and all intervals crossed zero.",
    "The result does not support a uniform short-lag conflict effect across these WHO critical-priority phenotypes in the available ATLAS surveillance sample.",
])

doc.add_heading("2. Wild-cluster and robustness analyses", level=1)
doc.add_paragraph(
    "Because the number of country clusters ranged from 28 to 46, the primary clustered t inference was checked using 9,999 Rademacher wild-cluster score replications. Wild-bootstrap p values ranged from 0.220 to 0.533 and did not alter the conclusion."
)
wild_rows = []
for row in main.itertuples(index=False):
    wild_rows.append([row.endpoint_id, f"{row.wild_bootstrap_p:.3f}",
                      f"{row.wild_bootstrap_or_ci_low:.3f}–{row.wild_bootstrap_or_ci_high:.3f}"])
add_table(doc, ["Endpoint", "Wild-cluster p", "Wild one-step OR interval"], wild_rows, [4.2, 4.0, 8.1], size=8.4)

doc.add_paragraph(
    "The direction and magnitude remained similar after removing isolates with unavailable covariates, raising the country-year threshold to 50, omitting countries without outcome variation, and removing composition adjustment. Same-year and two-year-lag specifications were also non-significant. The largest secondary signal was K. pneumoniae meropenem resistance at a two-year lag (OR 1.095, 95% CI 0.982–1.221; p = 0.100), which remains exploratory and does not satisfy confirmatory significance."
)

sens_rows = []
for endpoint in labels:
    for specification in ["Minimum cell n=50", "Same-year conflict", "Two-year lagged conflict", "Exclude countries without outcome variation"]:
        row = sensitivity.loc[(sensitivity.endpoint_id.eq(endpoint)) & (sensitivity.specification.eq(specification))].iloc[0]
        sens_rows.append([endpoint, specification, f"{row['or']:.3f} ({row.or_ci_low:.3f}–{row.or_ci_high:.3f})", f"{row.p_value:.3f}"])
add_table(doc, ["Endpoint", "Specification", "OR (95% CI)", "p value"], sens_rows, [2.3, 6.3, 4.5, 2.1], size=7.4)

doc.add_heading("3. Standardised AMR estimates", level=1)
doc.add_paragraph(
    "Country-year resistance was estimated using a logistic mixed model with a partially pooled country-year intercept and fixed effects for calendar year, age group, sex, specimen source and specialty. Each estimate was marginalised to the pooled endpoint-specific patient/specimen distribution. These are standardised resistance estimates among ATLAS surveillance isolates, not national prevalence estimates."
)
doc.add_picture(str(trend_path), width=Inches(6.25))
p = doc.add_paragraph("Figure 2. Isolate-weighted standardised and pooled crude resistance among contributing locked country-year cells. Country composition varies by year; the lines are descriptive.")
p.style = doc.styles["Caption"]

trend_rows = []
for endpoint in labels:
    a = trends.loc[(trends.endpoint_id.eq(endpoint)) & (trends.year.eq(2019))].iloc[0]
    b = trends.loc[(trends.endpoint_id.eq(endpoint)) & (trends.year.eq(2024))].iloc[0]
    cal = calibration.loc[calibration.endpoint_id.eq(endpoint)].iloc[0]
    trend_rows.append([endpoint, f"{a.isolate_weighted_standardised_pct:.1f}%", f"{b.isolate_weighted_standardised_pct:.1f}%",
                       f"{cal.mean_absolute_adjustment_pp:.2f}", f"{cal.crude_standardised_correlation:.3f}"])
add_table(doc, ["Endpoint", "2019 standardised", "2024 standardised", "Mean absolute adjustment, pp", "Crude-standardised r"],
          trend_rows, [2.5, 3.0, 3.0, 4.0, 3.7], size=8.0)
doc.add_paragraph(
    "All four mixed models converged. Standardisation altered individual country-year estimates by a mean of 1.18–2.76 percentage points, while crude-standardised correlations remained at least 0.993. The apparent decline from 2019 to 2024 is descriptive because the endpoint-specific panel is unbalanced and the contributing country set changes over time."
)

doc.add_heading("Exploratory escalation trajectories", level=2)
doc.add_paragraph(
    "The 17 Phase 3 escalation windows across Ivory Coast, Greece, Poland, Spain and Ukraine were summarised descriptively. Because only five countries contribute and pre/post coverage is irregular, no pooled hypothesis test or causal event-study coefficient was fitted. Crude resistance changes were heterogeneous."
)
event_rows = []
for endpoint, group in events.groupby("endpoint_id"):
    event_rows.append([endpoint, len(group), f"{group.index_minus_pre_pp.mean():+.1f}",
                       f"{group.post_minus_pre_pp.mean():+.1f}",
                       ", ".join(sorted(group.country.unique()))])
add_table(doc, ["Endpoint", "Windows", "Mean index−pre, pp", "Mean post−pre, pp", "Countries"],
          event_rows, [2.5, 1.7, 3.1, 3.1, 6.0], size=7.8)

doc.add_page_break()
doc.add_heading("4. MIC sensitivity analysis", level=1)
doc.add_paragraph(
    "MIC inequalities were not converted to exact values in the data pipeline. For a transparent sensitivity analysis, country-year mean log2 MIC was calculated under two bounded substitutions: at the observed assay boundary and one two-fold dilution beyond an open boundary. Weighted fixed-effects models used country-clustered inference. This is a sensitivity analysis, not a full interval-censored likelihood model."
)
mic_rows = []
for row in mic.itertuples(index=False):
    mic_rows.append([row.endpoint_id, row.specification, f"{row.log2_mic_change_per_exposure_doubling:+.3f}",
                     f"{row.ci_low:+.3f} to {row.ci_high:+.3f}", f"{row.p_value:.3f}"])
add_table(doc, ["Endpoint", "Censoring assumption", "log2 MIC change", "95% CI", "p value"],
          mic_rows, [2.2, 5.4, 2.7, 3.7, 1.8], size=7.7)
doc.add_paragraph("No MIC sensitivity coefficient was statistically significant, and the boundary assumptions did not materially change direction or magnitude.")

doc.add_heading("5. Diagnostics and limitations", level=1)
diag_rows = []
for row in diagnostics.itertuples(index=False):
    diag_rows.append([row.endpoint_id, "Yes" if row.converged else "No", row.clusters,
                      f"{row.pearson_dispersion:.3f}", row.most_influential_iso3,
                      f"{row.leave_one_cluster_or_min:.3f}–{row.leave_one_cluster_or_max:.3f}"])
add_table(doc, ["Endpoint", "Converged", "Clusters", "Pearson dispersion", "Largest influence", "Leave-one-country OR range"],
          diag_rows, [2.3, 2.0, 1.7, 3.0, 3.0, 4.2], size=8.0)
add_bullets(doc, [
    "Pearson dispersion was close to 1 for all four models, and leave-one-country estimates did not reverse the overall conclusion.",
    "Ireland and New Zealand had no meropenem-resistant K. pneumoniae isolates in their eligible records. Excluding both left the exposure estimate unchanged to three decimals.",
    "Some country fixed effects produced probabilities near 0 or 1, reflecting highly homogeneous country-level outcomes. The exposure coefficients nevertheless converged and were stable in influence and outcome-variation sensitivities.",
    "ATLAS is surveillance sampling rather than population-representative national surveillance. Conflict may also change who is cultured, which facilities report and which countries remain observed.",
    "The analysis estimates a within-country temporal association and cannot by itself establish causality. Residual time-varying confounding, exposure measurement error and short follow-up remain possible.",
    "The distributed-lag and event-study evidence is limited. Absence of a one-year association does not exclude effects operating through other delays, local subnational shocks or mechanisms not captured by annual national event counts.",
])

doc.add_heading("6. Phase 4 lock and next use", level=1)
add_note(doc, "Locked conclusion", "The Phase 4 confirmatory analysis does not detect an association between preceding-year national political-violence event intensity and the four prespecified AMR outcomes. All point estimates, confidence intervals, adjusted p values and null results must be retained in subsequent reporting.", AMBER)
add_bullets(doc, [
    "Carry the standardised country-year estimates and intervals into the dashboard, with n < 30 suppression retained.",
    "Treat event studies, alternative lags and MIC results as exploratory or sensitivity analyses.",
    "Proceed to Phase 5 without redefining the primary outcome family. One Health models should be labelled secondary and should not be used to rescue a null primary hypothesis.",
    "For publication, emphasise precision, uncertainty, surveillance limitations and heterogeneity rather than a binary significant/non-significant narrative.",
])

doc.add_heading("Reproducibility record", level=1)
doc.add_paragraph(
    "The analysis was generated by repository pipeline version 0.4.0 from the Phase 3 locked sample. The repository contains source code, frozen configuration, automated tests and machine-readable result tables. Raw and isolate-level Vivli data remain excluded from the distributable repository. All 12 automated tests passed."
)

path = OUT / "Phase4_Statistical_Analysis_Report.docx"
doc.save(path)
print(path)
