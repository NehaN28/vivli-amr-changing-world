from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path("/workspace/scratch/6f657a5e2035")
REPO = ROOT / "amr-changing-world"
AUDIT = REPO / "outputs" / "phase5_analysis"
OUT = ROOT / "outputs" / "phase5"
FIG = OUT / "figures"
OUT.mkdir(parents=True, exist_ok=True)
FIG.mkdir(parents=True, exist_ok=True)

oh = pd.read_csv(AUDIT / "one_health_models.csv")
woah = pd.read_csv(AUDIT / "woah_models.csv")
pathogen = pd.read_csv(AUDIT / "rd_pathogen_alignment.csv")
sector = pd.read_csv(AUDIT / "rd_sector_portfolio.csv")
area = pd.read_csv(AUDIT / "rd_research_area_portfolio.csv")
cross = pd.read_csv(AUDIT / "rd_cross_sector_portfolio.csv")
profile = pd.read_csv(AUDIT / "country_context_profile.csv")

labels = {
    "ECO_CAZ_R": r"$\it{E.\ coli}$–ceftazidime",
    "KPN_CAZ_R": r"$\it{K.\ pneumoniae}$–ceftazidime",
    "KPN_MEM_R": r"$\it{K.\ pneumoniae}$–meropenem",
    "ABA_MEM_R": r"$\it{A.\ baumannii}$–meropenem",
}

plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 8})
main = oh.loc[oh.analysis_family.eq("determinant_main_effect")].copy()
main["label"] = main["analysis"].str.replace("Lagged ", "", regex=False) + " | " + main["endpoint_id"].map(labels)
main = main.sort_values(["analysis", "endpoint_id"]).reset_index(drop=True)
fig, ax = plt.subplots(figsize=(7.2, 6.2))
y = np.arange(len(main))
ax.errorbar(main.odds_ratio, y, xerr=[main.odds_ratio-main.or_ci_low, main.or_ci_high-main.odds_ratio],
            fmt="o", color="black", ecolor="#777777", capsize=2, markersize=3)
ax.axvline(1, color="#777777", linestyle="--", linewidth=1)
ax.set_yticks(y, main.label)
ax.set_xscale("log")
ax.set_xlabel("Adjusted odds ratio (95% CI); scales differ by determinant")
ax.grid(axis="x", color="#E5E5E5", linewidth=.5)
ax.spines[["top", "right"]].set_visible(False)
fig.tight_layout()
forest = FIG / "Figure_1_One_Health_Main_Effects.png"
fig.savefig(forest, dpi=300, bbox_inches="tight")
plt.close(fig)

plot_pathogen = pathogen.loc[~pathogen.pathogen.eq("Not pathogen-specific")].sort_values("fractional_funding_usd")
fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.6))
axes[0].barh(plot_pathogen.pathogen, plot_pathogen.fractional_funding_usd / 1e9, color="#555555")
axes[0].set_xlabel("2015–2024 funding (USD billions)")
axes[0].set_title("Pathogen-specific allocation")
plot_sector = sector.sort_values("fractional_funding_usd")
axes[1].barh(plot_sector.sector, plot_sector.funding_share_pct, color="#777777")
axes[1].set_xlabel("Share of funding (%)")
axes[1].set_title("Sector allocation")
for ax in axes:
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="x", color="#E5E5E5", linewidth=.5)
fig.tight_layout()
rdfig = FIG / "Figure_2_RD_Portfolio.png"
fig.savefig(rdfig, dpi=300, bbox_inches="tight")
plt.close(fig)

NAVY = "275D7A"; PALE = "DDEBF3"; GRAY = "F2F4F5"; GREEN = "D9EAD3"; AMBER = "FCE8B2"; BORDER = "CCD4D9"


def shade(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); cell._tc.get_or_add_tcPr().append(shd)


def cell_text(cell, value, bold=False, color=None, size=8, center=False):
    cell.text = ""; p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0); r = p.add_run(str(value)); r.bold = bold; r.font.name = "Calibri"; r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None, size=8):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for i, h in enumerate(headers): cell_text(t.rows[0].cells[i], h, True, "FFFFFF", size, True); shade(t.rows[0].cells[i], NAVY)
    for vals in rows:
        cells = t.add_row().cells
        for i, v in enumerate(vals): cell_text(cells[i], v, False, None, size, i > 0)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def note(doc, label, text, fill=PALE):
    t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
    cell_text(t.cell(0, 0), label, True, NAVY, 9); cell_text(t.cell(0, 1), text, False, None, 9)
    shade(t.cell(0, 0), fill); shade(t.cell(0, 1), fill)


def bullets(doc, items):
    for x in items:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(x)


doc = Document(); sec = doc.sections[0]
sec.top_margin = Inches(.75); sec.bottom_margin = Inches(.75); sec.left_margin = Inches(.75); sec.right_margin = Inches(.75)
styles = doc.styles
styles["Normal"].font.name = "Calibri"; styles["Normal"].font.size = Pt(9); styles["Normal"].paragraph_format.space_after = Pt(5)
for name, size in [("Title", 22), ("Heading 1", 15), ("Heading 2", 11)]:
    styles[name].font.name = "Calibri"; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(80)
r = p.add_run("Phase 5\nOne Health Determinants and AMR R&D Alignment"); r.bold = True; r.font.size = Pt(23); r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Human AMR trends in a changing world\nScientific analysis report | Repository version 0.5.0\n23 July 2026")
doc.add_paragraph().add_run("Scope").bold = True
doc.add_paragraph("Secondary and exploratory analyses using the locked Phase 3 sample and Phase 4 standardised AMR estimates. No Phase 1 confirmatory hypothesis was altered.")
note(doc, "Bottom line", "No temperature, livestock-size, animal-AMU, or conflict-modification association survived false-discovery-rate correction, except inverse exploratory associations between swine livestock-unit share and both K. pneumoniae outcomes. These compositional ecological associations should not be interpreted as protective or causal.", GREEN)
doc.add_page_break()

doc.add_heading("1. Objectives and analysis hierarchy", level=1)
bullets(doc, [
    "Estimate associations of lagged temperature anomaly, livestock-system size and livestock composition with the four prespecified human AMR outcomes.",
    "Test whether temperature, livestock structure or animal antimicrobial use modifies the preceding-year conflict association.",
    "Analyse WOAH adjusted total animal AMU and a cephalosporin-class proxy in the restricted complete-case subset.",
    "Describe how 2015–2024 AMR R&D funding is distributed by pathogen, sector, research area and recipient institution country.",
    "Create a dashboard-ready country context profile without collapsing heterogeneous and incomplete components into a single vulnerability score.",
])
doc.add_heading("Model definitions", level=2)
table(doc, ["Component", "Definition", "Inference"], [
    ["Temperature", "Previous-year annual anomaly, per 1 °C", "Country/year fixed effects; country-clustered uncertainty"],
    ["Livestock size", "Previous-year total livestock units, per doubling", "Not density; no population or land denominator"],
    ["Livestock composition", "Cattle/buffalo, poultry or swine share, per 10 percentage points", "Separate models because shares are compositional"],
    ["Animal AMU", "Previous-year adjusted total mg/kg; log2(1 + mg/kg)", "WOAH complete cases only"],
    ["Aligned class", "All-generation cephalosporin mg/kg proxy for ceftazidime outcomes", "3rd/4th-generation exported series contained zero values"],
    ["Effect modification", "Conflict × centred One Health component", "Exploratory interaction term"],
], [1.15, 3.8, 2.2])
doc.add_paragraph("All isolate-level models adjust for year, country, age group, sex, specimen source and specialty. Conflict is retained as a covariate. Temperature is also retained in WOAH models. Benjamini–Hochberg false-discovery-rate correction is applied separately to each analysis family.")
doc.add_page_break()

doc.add_heading("2. One Health determinant results", level=1)
doc.add_picture(str(forest), width=Inches(6.6))
cap = doc.add_paragraph("Figure 1. Adjusted main-effect estimates. Scales differ: temperature per 1 °C, livestock units per doubling, and species shares per 10 percentage points.")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig = oh.loc[oh.fdr_significant_005]
rows = []
for x in sig.itertuples(index=False):
    rows.append([x.endpoint, x.analysis, f"{x.odds_ratio:.3f} ({x.or_ci_low:.3f}–{x.or_ci_high:.3f})", f"{x.p_value:.4f}", f"{x.fdr_p:.4f}"])
table(doc, ["Outcome", "Exposure", "Adjusted OR (95% CI)", "p", "FDR p"], rows, [1.7, 2.2, 1.5, .7, .7])
bullets(doc, [
    "Swine livestock-unit share was inversely associated with K. pneumoniae ceftazidime resistance: OR 0.688 per 10-percentage-point increase (95% CI 0.551–0.859; FDR p = 0.0149).",
    "Swine share was also inversely associated with K. pneumoniae meropenem resistance: OR 0.458 (95% CI 0.306–0.686; FDR p = 0.0068).",
    "No temperature main effect or conflict × temperature interaction survived correction.",
    "No livestock-size or livestock-composition conflict interaction survived correction. The smallest interaction FDR p was 0.369.",
])
note(doc, "Interpretation", "The swine-share findings are hypothesis-generating. A larger share necessarily means a smaller share of other livestock groups, and within-country changes may reflect economic, reporting or agricultural transitions. No causal or protective claim is warranted.", AMBER)
doc.add_page_break()

doc.add_heading("3. WOAH animal antimicrobial-use analysis", level=1)
coverage = woah.groupby("endpoint_id").agg(countries=("countries", "max"), cells=("country_year_cells", "max"), isolates=("tested_isolates", "max")).reset_index()
plain_labels = {
    "ECO_CAZ_R": "E. coli–ceftazidime", "KPN_CAZ_R": "K. pneumoniae–ceftazidime",
    "KPN_MEM_R": "K. pneumoniae–meropenem", "ABA_MEM_R": "A. baumannii–meropenem",
}
table(doc, ["Outcome", "Countries", "Country-years", "Isolates"], [[plain_labels[x.endpoint_id], int(x.countries), int(x.cells), f"{int(x.isolates):,}"] for x in coverage.itertuples(index=False)], [3.2, 1, 1.2, 1.2])
doc.add_paragraph("No adjusted or unadjusted total animal-AMU association, cephalosporin-proxy association, or conflict interaction survived FDR correction. The smallest corrected p values were 0.331 for AMU main effects and 0.432 for AMU interactions.")
bullets(doc, [
    "Only 14 countries contributed to the E. coli subset, 13 to each K. pneumoniae subset and 6 to A. baumannii.",
    "Adjusted total AMU and unadjusted total AMU gave broadly similar, imprecise estimates.",
    "The all-generation cephalosporin proxy was used only for ceftazidime outcomes; no carbapenem-specific WOAH class was available.",
    "These are limited-coverage ecological analyses. Non-significance should be described as insufficient evidence, not absence of an animal–human AMR relationship.",
])
note(doc, "Decision", "WOAH will remain a secondary dashboard layer and supplementary analysis. It should not be used to restrict the main global panel or to construct a complete global vulnerability score.")
doc.add_page_break()

doc.add_heading("4. R&D investment alignment", level=1)
doc.add_picture(str(rdfig), width=Inches(7.0))
cap = doc.add_paragraph("Figure 2. Fractionally allocated AMR R&D funding for projects starting in 2015–2024. Amounts are nominal reported USD award commitments.")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
human = sector.loc[sector.sector.eq("Human")].iloc[0]
animal = sector.loc[sector.sector.eq("Animal")].iloc[0]
environment = sector.loc[sector.sector.eq("Environment")].iloc[0]
cross_row = cross.loc[cross.one_health_cross_sector.eq(True)].iloc[0]
bullets(doc, [
    f"Human-sector work accounted for {human.funding_share_pct:.1f}% of fractionally allocated funding, compared with {animal.funding_share_pct:.1f}% for animal and {environment.funding_share_pct:.1f}% for environment.",
    f"Only {cross_row.funding_share_pct:.1f}% of reported funding was attached to projects coded across more than one sector.",
    f"Basic research represented {area.iloc[0].funding_share_pct:.1f}% of funding, followed by therapeutics ({area.loc[area.research_area.eq('Therapeutics'),'funding_share_pct'].iloc[0]:.1f}%) and operational research ({area.loc[area.research_area.eq('Operational'),'funding_share_pct'].iloc[0]:.1f}%).",
    "Among mapped pathogen-specific funding, S. aureus received 26.1%, E. coli 24.3%, P. aeruginosa 22.7%, K. pneumoniae 9.2% and A. baumannii 8.4%.",
])
align = pathogen.loc[pathogen.pathogen.isin(["Escherichia coli", "Klebsiella pneumoniae", "Acinetobacter baumannii"])]
table(doc, ["Pathogen", "Specific funding share", "2024 standardised resistance", "2019–2024 change"], [[x.pathogen, f"{x.share_of_pathogen_specific_funding_pct:.1f}%", f"{x.latest_standardised_resistance_pct:.1f}%", f"{x.annual_change_pp:+.2f} pp/year"] for x in align.itertuples(index=False)], [2.1, 1.4, 1.7, 1.6])
note(doc, "Policy signal", "A. baumannii had the highest observed standardised resistance among the three directly comparable primary pathogens, but a smaller pathogen-specific funding share. This is an attention-alignment signal, not proof of underfunding: resistance percentage is not disease burden, investments may precede the analysis window, and project benefit is not confined to the recipient institution country.", AMBER)
doc.add_page_break()

doc.add_heading("5. Vulnerability framework for the dashboard", level=1)
doc.add_paragraph("The dashboard dataset contains separate country-level components rather than a single score. This preserves interpretability and avoids converting missing WOAH data or null interaction estimates into an artificial ranking.")
table(doc, ["Component", "Dashboard measure", "Caution"], [
    ["Human AMR level", "Mean 2024 standardised resistance across available primary endpoints", "ATLAS isolates, not national prevalence"],
    ["Human AMR trajectory", "Mean within-country change in percentage points per year", "Requires at least three years"],
    ["Conflict", "2023 events and log2(1 + events)", "Exposure intensity, not individual experience"],
    ["Temperature", "2023 anomaly in °C and baseline-SD units", "Country annual average"],
    ["Livestock", "Total units and species shares", "Not density without denominator"],
    ["Animal AMU", "Adjusted mg/kg where reported", "Substantial missingness"],
    ["R&D capacity", "2015–2024 recipient projects and reported commitments", "Recipient location is not study location or benefit"],
    ["Reliability", "Endpoints, years, isolates and component availability", "Shown beside every profile"],
], [1.25, 3.2, 2.3])
doc.add_paragraph(f"The Phase 5 profile contains {len(profile)} countries with 2024 standardised AMR data. Each row records component availability and explicitly sets composite_score_constructed = FALSE.")
doc.add_heading("Overall conclusion", level=2)
bullets(doc, [
    "The Phase 4 null conflict result remains unchanged.",
    "One Health effect-modification models did not identify a reproducible vulnerability modifier after correction.",
    "The swine-share associations merit sensitivity analysis with population or agricultural-land denominators and additional economic covariates before publication-level emphasis.",
    "The clearest policy finding is the strong human-sector concentration of AMR R&D and limited cross-sector funding.",
    "Phase 6 should display evidence strength and missingness alongside every visual, and should not imply causal effects or a validated global vulnerability rank.",
])
doc.add_page_break()

doc.add_heading("6. Limitations and reproducibility decisions", level=1)
bullets(doc, [
    "ATLAS is a surveillance isolate collection. Standardised resistance estimates are not population-representative national prevalence.",
    "Country and year fixed effects reduce time-invariant country confounding but do not remove time-varying confounding or reverse causation.",
    "Livestock units use documented analytical conversion weights. They are not measured biomass, and density could not be derived without population or land denominators.",
    "Livestock shares are compositional. Separate models do not fully solve substitution interpretation.",
    "WOAH reporting is incomplete and may differ by reporting option. Class-specific data do not include a carbapenem class and the narrow 3rd/4th-generation cephalosporin export was unusable.",
    "R&D funding amounts are nominal reported commitments, not inflation-adjusted expenditure. Keyword-derived pathogen tags and fractional allocation reduce double counting but can misclassify broad projects.",
    "Recipient institution geography describes research-capacity concentration, not where research was conducted or where benefits accrue.",
    "All Phase 5 analyses are secondary or exploratory. FDR correction was performed within prespecified analysis families.",
])
doc.add_heading("Reproducible outputs", level=2)
doc.add_paragraph("Repository version 0.5.0 contains the source transformations, model code, tests, CSV result tables and generation scripts for this report and the accompanying workbook. Restricted isolate-level Vivli data remain excluded from the public repository.")
doc.add_heading("Data-source references", level=2)
doc.add_paragraph("WHO Bacterial Priority Pathogens List 2024; ATLAS data accessed through Vivli; ACLED political-violence event counts; FAOSTAT Temperature Change and livestock stocks; WOAH antimicrobial-use data; Global AMR R&D Hub project export dated 31 May 2026. Exact source filenames and hashes are recorded in the repository manifest.")

footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Phase 5 | Human AMR trends in a changing world | v0.5.0")
out = OUT / "Phase5_One_Health_and_RD_Analysis_Report.docx"
doc.save(out)
print(out)
